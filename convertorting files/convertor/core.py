"""
File Conversion Module

This module provides functionality for extracting content from various file formats
and processing it using AI for summarization or description.

Supported file formats:
- Text Documents:
  - Plain Text (.txt)
  - PDF Documents (.pdf)
  - Microsoft Word Documents (.docx, .doc)
  - OpenDocument Text (.odt)
  - HTML Documents (.html, .htm)
- Structured Data:
  - Excel Spreadsheets (.xlsx, .xls)
  - CSV Files (.csv)
  - JSON Files (.json)
- Images:
  - PNG Images (.png)
  - JPEG Images (.jpg, .jpeg)
  - WebP Images (.webp)
  - GIF Images (.gif)
  - BMP Images (.bmp)

The module uses Google's Gemini API for AI processing, which requires an API key
to be set in the GEMINI_API_KEY environment variable.
"""

import os
import google.generativeai as genai
import logging
import subprocess # Added for running external commands like antiword
import json
import csv
import tempfile
from pathlib import Path
from PIL import Image # Import Pillow
import PyPDF2 # Import PyPDF2 for PDF parsing
import docx # Import python-docx for DOCX parsing
import pandas as pd # For Excel and CSV processing
import openpyxl # For direct Excel processing
from bs4 import BeautifulSoup # For HTML parsing
import odf.opendocument # For ODT parsing
from odf.text import P, Span # For ODT text elements
import requests # For fetching HTML from URLs
from .email_parser import is_email_correspondence, process_email_correspondence # Import email parser

# Import OCR libraries (with fallback if not installed)
try:
    import pytesseract
    import pdf2image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR libraries (pytesseract/pdf2image) not available. OCR functionality will be disabled.")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Gemini API Configuration ---
# Ensure GEMINI_API_KEY environment variable is set
try:
    gemini_api_key = os.environ.get("GEMINI_API_KEY")
    if not gemini_api_key:
        logging.warning("GEMINI_API_KEY environment variable not set. AI processing will fail.")
        # raise ValueError("GEMINI_API_KEY environment variable not set.")
    else:
        genai.configure(api_key=gemini_api_key)
except Exception as e:
    logging.error(f"Error configuring Gemini API: {e}")
    gemini_api_key = None # Ensure it's None if config fails

# --- File Parsing Functions (Placeholders) ---

def parse_text_file(file_path):
    """Parses plain text files."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logging.error(f"Error parsing text file {file_path}: {e}")
        raise

def extract_text_from_pdf_with_ocr(file_path, max_pages=3, dpi=200):
    """
    Extracts text from PDF using OCR (Optical Character Recognition).

    This function is used as a fallback when regular text extraction fails.
    It converts PDF pages to images and then uses pytesseract to extract text.

    Default max_pages is set to 3 to avoid processing too many pages at once,
    which can cause performance issues or timeouts.

    Args:
        file_path (str): Path to the PDF file
        max_pages (int, optional): Maximum number of pages to process. Defaults to 3 pages.
        dpi (int, optional): DPI for image conversion. Defaults to 200 for balance of quality and speed.

    Returns:
        str: Extracted text content from the PDF file

    Raises:
        ImportError: If OCR libraries are not available
        Exception: For other unexpected errors
    """
    if not OCR_AVAILABLE:
        logging.error("OCR libraries not available. Cannot perform OCR on PDF.")
        raise ImportError("OCR libraries (pytesseract/pdf2image) not installed. Please install them to use OCR functionality.")

    text_parts = []

    try:
        logging.info(f"Converting PDF to images for OCR: {os.path.basename(file_path)}")

        # Create a temporary directory for the images
        with tempfile.TemporaryDirectory() as temp_dir:
            # Convert PDF to images
            images = pdf2image.convert_from_path(
                file_path,
                dpi=dpi,
                output_folder=temp_dir,
                fmt="png",
                thread_count=4  # Use multiple threads for faster conversion
            )

            # Limit the number of pages if specified
            if max_pages:
                images = images[:max_pages]

            total_pages = len(images)
            logging.info(f"Processing {total_pages} pages with OCR")

            # Process each image with OCR
            for i, img in enumerate(images):
                try:
                    logging.info(f"Performing OCR on page {i+1}/{total_pages}")

                    # Use pytesseract to extract text from the image
                    page_text = pytesseract.image_to_string(img, lang='ces+eng')  # Use Czech and English languages

                    if page_text.strip():
                        text_parts.append(page_text)
                    else:
                        logging.warning(f"OCR could not extract text from page {i+1}")
                except Exception as ocr_error:
                    logging.error(f"OCR error on page {i+1}: {ocr_error}")

        # Join all text parts with newlines
        result = "\n".join(text_parts)

        if not result.strip():
            logging.warning("OCR did not extract any text from the PDF")
            return None

        logging.info(f"Successfully extracted {len(result)} characters with OCR")
        return result

    except Exception as e:
        logging.error(f"Error during OCR processing: {e}")
        raise

def parse_pdf_file(file_path, max_pages=5, page_batch_size=5, try_ocr=True):
    """
    Parses PDF files using PyPDF2 with OCR fallback.

    Optimized for large files with batch processing and progress logging.
    If regular text extraction fails and try_ocr is True, falls back to OCR.

    Default max_pages is set to 5 to avoid processing too many pages at once,
    which can cause API rate limits or timeouts.

    Args:
        file_path (str): Path to the PDF file
        max_pages (int, optional): Maximum number of pages to process. Defaults to 5 pages.
        page_batch_size (int, optional): Number of pages to process in each batch. Defaults to 5.
        try_ocr (bool, optional): Whether to try OCR if regular text extraction fails. Defaults to True.

    Returns:
        str: Extracted text content from the PDF file

    Raises:
        ValueError: If the file is not a valid PDF
        IOError: If there's an issue reading the file
        Exception: For other unexpected errors
    """
    text_parts = []  # Store text in parts to avoid large string concatenation
    try:
        # Check file size before processing
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)

        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            total_pages = len(reader.pages)

            # Determine how many pages to process
            pages_to_process = min(total_pages, max_pages) if max_pages else total_pages

            logging.info(f"Parsing PDF: {os.path.basename(file_path)} ({pages_to_process} of {total_pages} pages, {file_size_mb:.2f}MB)")

            # Process pages in batches to avoid memory issues with large PDFs
            for batch_start in range(0, pages_to_process, page_batch_size):
                batch_end = min(batch_start + page_batch_size, pages_to_process)
                logging.info(f"Processing PDF pages {batch_start+1}-{batch_end} of {pages_to_process}")

                for i in range(batch_start, batch_end):
                    try:
                        page = reader.pages[i]
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                        else:
                            logging.warning(f"Could not extract text from page {i+1} of {os.path.basename(file_path)}")
                    except Exception as page_e:
                        logging.error(f"Error extracting text from page {i+1} of {file_path}: {page_e}")

            # Check if we processed all pages or limited the number
            if total_pages > pages_to_process:
                logging.info(f"Processed {pages_to_process} pages out of {total_pages} total pages in {os.path.basename(file_path)}")
                text_parts.append(f"\n[Note: Only {pages_to_process} of {total_pages} pages were processed due to size limits]")

            logging.info(f"Successfully parsed PDF: {os.path.basename(file_path)}")

        # Join all text parts with newlines
        result = "\n".join(text_parts)

        # Check if we extracted any meaningful text
        if not result.strip() and try_ocr and OCR_AVAILABLE:
            logging.info(f"No text extracted from PDF using PyPDF2, trying OCR: {os.path.basename(file_path)}")
            try:
                ocr_result = extract_text_from_pdf_with_ocr(file_path, max_pages)
                if ocr_result:
                    logging.info(f"Successfully extracted text using OCR: {os.path.basename(file_path)}")
                    return ocr_result
                else:
                    logging.warning(f"OCR also failed to extract text from PDF: {os.path.basename(file_path)}")
            except ImportError as ie:
                logging.warning(f"OCR libraries not available: {ie}")
            except Exception as ocr_e:
                logging.error(f"Error during OCR processing: {ocr_e}")

        # For very large content, provide a summary of the extraction
        if len(result) > 1000000:  # If more than ~1MB of text
            chars_per_page = len(result) / pages_to_process if pages_to_process > 0 else 0
            logging.info(f"Extracted {len(result)} characters from PDF ({chars_per_page:.1f} chars/page)")

        return result
    except PyPDF2.errors.PdfReadError as pdf_err:
        logging.error(f"Invalid PDF file {file_path}: {pdf_err}")

        # Try OCR as a fallback for corrupted PDFs
        if try_ocr and OCR_AVAILABLE:
            logging.info(f"PDF reading error, trying OCR as fallback: {os.path.basename(file_path)}")
            try:
                ocr_result = extract_text_from_pdf_with_ocr(file_path, max_pages)
                if ocr_result:
                    logging.info(f"Successfully extracted text using OCR fallback: {os.path.basename(file_path)}")
                    return ocr_result
            except Exception as ocr_e:
                logging.error(f"OCR fallback also failed: {ocr_e}")

        raise ValueError(f"Invalid or corrupted PDF file: {os.path.basename(file_path)}") from pdf_err
    except IOError as e:
        logging.error(f"IO error reading PDF file {file_path}: {e}")
        raise IOError(f"Error reading file: {e}")
    except Exception as e:
        logging.error(f"Error parsing PDF file {file_path}: {e}")
        raise  # Re-raise to be caught by process_file


def parse_docx_file(file_path, include_tables=True, batch_size=100, max_pages=None):
    """
    Parses DOCX files using python-docx.

    Optimized for large files with batch processing and progress logging.
    Includes table parsing by default.

    Args:
        file_path (str): Path to the DOCX file
        include_tables (bool, optional): Whether to include tables in the extracted text. Defaults to True.
        batch_size (int, optional): Number of paragraphs to process in each batch. Defaults to 100.
        max_pages (int, optional): Maximum number of pages to process. If None, process all pages.

    Returns:
        str: Extracted text content from the DOCX file

    Raises:
        ValueError: If the file is not a valid DOCX
        IOError: If there's an issue reading the file
        Exception: For other unexpected errors
    """
    text_parts = []  # Store text in parts to avoid large string concatenation
    try:
        # Check file size before processing
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        logging.info(f"Parsing DOCX: {os.path.basename(file_path)} ({file_size_mb:.2f}MB)")

        try:
            doc = docx.Document(file_path)
        except Exception as doc_error:
            logging.error(f"Failed to open DOCX file {file_path}: {doc_error}")
            raise ValueError(f"Invalid DOCX format: {doc_error}")

        # Process paragraphs in batches
        total_paragraphs = len(doc.paragraphs)
        logging.info(f"DOCX contains {total_paragraphs} paragraphs")

        # Odhadneme počet stránek na základě počtu odstavců (přibližně 20 odstavců na stránku)
        estimated_pages = total_paragraphs // 20 + 1
        logging.info(f"Estimated pages in DOCX: {estimated_pages}")

        # Určíme, kolik odstavců zpracovat na základě max_pages
        paragraphs_to_process = total_paragraphs
        if max_pages is not None and estimated_pages > max_pages:
            paragraphs_to_process = max_pages * 20  # Přibližný počet odstavců pro max_pages
            logging.info(f"Limiting processing to approximately {max_pages} pages ({paragraphs_to_process} paragraphs)")

        # Zpracujeme všechny odstavce, pokud se jedná o emailovou korespondenci
        # Emailová korespondence vyžaduje zpracování celého dokumentu pro správnou detekci
        # Zkontrolujeme, zda dokument obsahuje typické znaky emailové korespondence
        email_indicators = ['Od:', 'Odesílatel:', 'From:', 'Komu:', 'To:', 'Předmět:', 'Subject:']
        email_indicator_count = 0

        # Zkontrolujeme prvních 50 odstavců (nebo méně, pokud jich je méně)
        for i in range(min(50, total_paragraphs)):
            para_text = doc.paragraphs[i].text
            for indicator in email_indicators:
                if indicator in para_text:
                    email_indicator_count += 1

        # Pokud najdeme alespoň 3 indikátory, považujeme to za emailovou korespondenci
        is_likely_email = email_indicator_count >= 3

        if is_likely_email:
            logging.info(f"Document appears to be email correspondence, processing all paragraphs")
            paragraphs_to_process = total_paragraphs

        for batch_start in range(0, paragraphs_to_process, batch_size):
            batch_end = min(batch_start + batch_size, paragraphs_to_process)
            if total_paragraphs > 1000:  # Only log batches for large documents
                logging.info(f"Processing DOCX paragraphs {batch_start+1}-{batch_end} of {paragraphs_to_process}")

            # Process this batch of paragraphs
            for i in range(batch_start, batch_end):
                try:
                    para = doc.paragraphs[i]
                    if para.text.strip():  # Only add non-empty paragraphs
                        text_parts.append(para.text)
                except Exception as para_error:
                    logging.error(f"Error extracting text from paragraph {i+1}: {para_error}")

        # Process tables if requested
        if include_tables and doc.tables:
            table_count = len(doc.tables)
            logging.info(f"Processing {table_count} tables in DOCX")

            for table_index, table in enumerate(doc.tables):
                try:
                    # Add a header for the table
                    text_parts.append(f"\nTable {table_index + 1}:")

                    # Process each row
                    for row in table.rows:
                        row_texts = []
                        for cell in row.cells:
                            row_texts.append(cell.text.strip())

                        # Join cells with tabs and add to text parts
                        if any(row_texts):  # Only add non-empty rows
                            text_parts.append("\t".join(row_texts))
                except Exception as table_error:
                    logging.error(f"Error extracting text from table {table_index + 1}: {table_error}")

        # Join all text parts with newlines
        result = "\n".join(text_parts)

        # For very large content, provide a summary of the extraction
        if len(result) > 1000000:  # If more than ~1MB of text
            logging.info(f"Extracted {len(result)} characters from DOCX")

        # Pokud jsme omezili počet stránek a není to emailová korespondence, přidáme poznámku
        if max_pages is not None and estimated_pages > max_pages and not is_likely_email:
            result += f"\n\n[Poznámka: Dokument obsahuje přibližně {estimated_pages} stránek, ale bylo zpracováno pouze přibližně {max_pages} stránek.]"

        logging.info(f"Successfully parsed DOCX: {os.path.basename(file_path)}")
        return result
    except IOError as e:
        logging.error(f"IO error reading DOCX file {file_path}: {e}")
        raise IOError(f"Error reading file: {e}")
    except ValueError as e:
        # Re-raise ValueError (invalid format)
        logging.error(f"Invalid DOCX format in file {file_path}: {e}")
        raise
    except Exception as e:
        # Catch potential errors like file corruption or unsupported formats within DOCX
        logging.error(f"Error parsing DOCX file {file_path}: {e}")
        raise  # Re-raise to be caught by process_file

def parse_doc_file(file_path):
    """
    Parses DOC files using antiword if available, otherwise provides a helpful error message.

    This function attempts to use the antiword command-line tool to extract text from .doc files.
    If antiword is not installed, it returns a helpful message instead of raising an exception.

    Args:
        file_path (str): Path to the DOC file

    Returns:
        str: Extracted text content from the DOC file or an error message

    Note:
        Antiword is required for full .doc file support. If not installed, a message
        will be returned instructing the user how to install it.
    """
    text = ""
    try:
        logging.info(f"Parsing DOC with antiword: {os.path.basename(file_path)}")
        # Check if antiword is installed
        try:
            # Try to run antiword with --help to check if it's installed
            subprocess.run(['antiword', '--help'], capture_output=True, check=False)
            antiword_installed = True
        except FileNotFoundError:
            antiword_installed = False

        if not antiword_installed:
            # Antiword not installed - return a helpful message instead of raising an exception
            install_message = """
Antiword is not installed. To process .doc files, please install antiword:

- On macOS: brew install antiword
- On Ubuntu/Debian: apt-get install antiword
- On Windows: Download from http://www.winfield.demon.nl/

Alternatively, consider converting your .doc file to .docx format which doesn't require antiword.
"""
            logging.warning(f"Antiword not installed. Cannot process DOC file: {os.path.basename(file_path)}")
            return install_message

        # Execute antiword, capture stdout, decode as UTF-8, ignore errors
        # Use '-m UTF-8' for UTF-8 output, '-w 0' for unlimited line width
        process = subprocess.run(
            ['antiword', '-m', 'UTF-8', '-w', '0', file_path],
            capture_output=True,
            text=True, # Decodes stdout/stderr using default encoding (usually utf-8)
            check=False # Don't raise exception on non-zero exit code, check manually
        )

        if process.returncode == 0:
            text = process.stdout
            logging.info(f"Successfully parsed DOC with antiword: {os.path.basename(file_path)}")
        else:
            # Log stderr if antiword failed
            error_message = process.stderr.strip() if process.stderr else "Unknown antiword error"
            logging.error(f"Antiword failed for {file_path}. Return code: {process.returncode}. Error: {error_message}")
            # Return error message instead of raising exception
            return f"Error processing .doc file: {error_message}"

        return text
    except Exception as e:
        logging.error(f"Error parsing DOC file {file_path} with antiword: {e}")
        # Return error message instead of raising exception
        return f"Error processing .doc file: {str(e)}"

def parse_xlsx_file(file_path, max_rows_per_sheet=1000, max_sheets=None):
    """
    Parses Excel (.xlsx) files using pandas.

    Optimized for large files with row and sheet limits.

    Args:
        file_path (str): Path to the Excel file
        max_rows_per_sheet (int, optional): Maximum number of rows to process per sheet. Defaults to 1000.
        max_sheets (int, optional): Maximum number of sheets to process. If None, process all sheets.

    Returns:
        str: Formatted string representation of the Excel content

    Raises:
        ValueError: If the file is not a valid Excel file
        IOError: If there's an issue reading the file
        Exception: For other unexpected errors
    """
    try:
        # Check file size before processing
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        logging.info(f"Parsing Excel file: {os.path.basename(file_path)} ({file_size_mb:.2f}MB)")

        # Get sheet names first to determine structure
        try:
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
        except Exception as excel_error:
            logging.error(f"Failed to open Excel file {file_path}: {excel_error}")
            raise ValueError(f"Invalid Excel format: {excel_error}")

        # Determine how many sheets to process
        total_sheets = len(sheet_names)
        sheets_to_process = min(total_sheets, max_sheets) if max_sheets else total_sheets

        if sheets_to_process < total_sheets:
            logging.info(f"Processing {sheets_to_process} of {total_sheets} sheets in Excel file")

        # Initialize result text
        text_parts = []
        text_parts.append("Excel File Content:")

        # Add sheet information
        if total_sheets > 1:
            text_parts.append(f"File contains {total_sheets} sheets: {', '.join(sheet_names)}")
            if sheets_to_process < total_sheets:
                text_parts.append(f"Note: Only processing first {sheets_to_process} sheets due to size limits.")
            text_parts.append("")  # Empty line

        # Process each sheet up to the limit
        for sheet_index, sheet_name in enumerate(sheet_names[:sheets_to_process]):
            try:
                logging.info(f"Processing Excel sheet {sheet_index+1}/{sheets_to_process}: {sheet_name}")

                # Read the sheet with row limits for large files
                try:
                    # First try to get row count without loading all data
                    sheet_preview = pd.read_excel(file_path, sheet_name=sheet_name, nrows=5)
                    if sheet_preview.empty:
                        text_parts.append(f"Sheet: {sheet_name} (empty)")
                        continue

                    # For large files, use chunksize to read in batches
                    if file_size_mb > 10:  # For files larger than 10MB
                        # Read with nrows parameter to limit rows
                        sheet_df = pd.read_excel(file_path, sheet_name=sheet_name, nrows=max_rows_per_sheet)
                        total_rows = len(sheet_df)

                        # Check if we might have truncated rows
                        if total_rows >= max_rows_per_sheet:
                            logging.info(f"Sheet {sheet_name} has at least {total_rows} rows, may be truncated")
                            truncated = True
                        else:
                            truncated = False
                    else:
                        # For smaller files, read the whole sheet
                        sheet_df = pd.read_excel(file_path, sheet_name=sheet_name)
                        total_rows = len(sheet_df)

                        # Check if we need to truncate
                        if total_rows > max_rows_per_sheet:
                            logging.info(f"Truncating sheet {sheet_name} from {total_rows} to {max_rows_per_sheet} rows")
                            sheet_df = sheet_df.head(max_rows_per_sheet)
                            truncated = True
                        else:
                            truncated = False

                    # Add sheet header
                    text_parts.append(f"Sheet: {sheet_name} ({total_rows} rows, {len(sheet_df.columns)} columns)")

                    # Add column names
                    text_parts.append(f"Columns: {', '.join(sheet_df.columns)}")

                    # Format the data
                    if not sheet_df.empty:
                        # For very wide DataFrames, limit the output
                        if len(sheet_df.columns) > 10:
                            logging.info(f"Sheet {sheet_name} has {len(sheet_df.columns)} columns, showing first 10")
                            text_parts.append("Note: Showing only first 10 columns due to width limits.")
                            sheet_df = sheet_df.iloc[:, :10]

                        # Convert to string and add to text parts
                        text_parts.append(sheet_df.to_string())

                        # Add truncation note if needed
                        if truncated:
                            text_parts.append(f"Note: Only showing first {max_rows_per_sheet} rows of this sheet.")

                    text_parts.append("")  # Empty line for separation

                except Exception as sheet_error:
                    logging.error(f"Error reading sheet {sheet_name}: {sheet_error}")
                    text_parts.append(f"Sheet: {sheet_name} (Error: {sheet_error})")

            except Exception as e:
                logging.error(f"Error processing sheet {sheet_name}: {e}")
                text_parts.append(f"Error processing sheet {sheet_name}: {e}")

        # Join all parts with newlines
        result = "\n".join(text_parts)

        logging.info(f"Successfully parsed Excel file: {os.path.basename(file_path)}")
        return result
    except pd.errors.EmptyDataError:
        logging.warning(f"Excel file {os.path.basename(file_path)} is empty")
        return "Empty Excel file"
    except IOError as e:
        logging.error(f"IO error reading Excel file {file_path}: {e}")
        raise IOError(f"Error reading file: {e}")
    except ValueError as e:
        # Re-raise ValueError (invalid format)
        logging.error(f"Invalid Excel format in file {file_path}: {e}")
        raise
    except Exception as e:
        logging.error(f"Error parsing Excel file {file_path}: {e}")
        raise

def parse_csv_file(file_path):
    """
    Parses CSV files using pandas.

    Handles various CSV formats with different delimiters and encodings.
    Attempts multiple parsing strategies if the initial attempt fails.

    Args:
        file_path (str): Path to the CSV file

    Returns:
        str: Formatted string representation of the CSV content

    Raises:
        ValueError: If the file is not a valid CSV
        UnicodeError: If the file encoding cannot be determined
        IOError: If there's an issue reading the file
        Exception: For other unexpected errors
    """
    try:
        logging.info(f"Parsing CSV file: {os.path.basename(file_path)}")

        # Check if file is empty
        if os.path.getsize(file_path) == 0:
            logging.warning(f"CSV file {os.path.basename(file_path)} is empty")
            return "Empty CSV file"

        # List of encodings to try
        encodings = ['utf-8', 'latin1', 'iso-8859-1', 'cp1252']

        # Try to detect encoding and delimiter
        df = None
        last_error = None

        # First try: Use pandas with default settings
        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                logging.info(f"Successfully parsed CSV with {encoding} encoding")
                break
            except UnicodeDecodeError:
                continue
            except pd.errors.ParserError:
                # If parsing fails, it might be due to delimiter
                try:
                    # Try to detect delimiter by reading first few lines
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        sample = f.read(1024)
                        try:
                            dialect = csv.Sniffer().sniff(sample)
                            df = pd.read_csv(file_path, sep=dialect.delimiter, encoding=encoding)
                            logging.info(f"Successfully parsed CSV with {encoding} encoding and {dialect.delimiter} delimiter")
                            break
                        except csv.Error:
                            # Sniffer failed, try common delimiters
                            for delimiter in [',', ';', '\t', '|']:
                                try:
                                    df = pd.read_csv(file_path, sep=delimiter, encoding=encoding)
                                    logging.info(f"Successfully parsed CSV with {encoding} encoding and {delimiter} delimiter")
                                    break
                                except:
                                    continue
                            if df is not None:
                                break
                except Exception as e:
                    last_error = e
                    continue
            except Exception as e:
                last_error = e
                continue

        # If all attempts failed
        if df is None:
            if last_error:
                raise ValueError(f"Failed to parse CSV file: {last_error}")
            else:
                raise ValueError("Failed to parse CSV file: Unknown error")

        # Check if DataFrame is empty
        if df.empty:
            logging.warning(f"CSV file {os.path.basename(file_path)} appears to be empty or has no valid data")
            return "CSV file contains no valid data"

        # Convert DataFrame to string representation
        text = "CSV File Content:\n\n"

        # Add metadata about the file
        text += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
        text += f"Column names: {', '.join(df.columns)}\n\n"

        # For very large DataFrames, limit the output
        if len(df) > 100:
            text += "Note: Showing first 50 and last 50 rows of a large CSV file.\n\n"
            text += "First 50 rows:\n"
            text += df.head(50).to_string()
            text += "\n\n...\n\n"
            text += "Last 50 rows:\n"
            text += df.tail(50).to_string()
        else:
            text += df.to_string()

        logging.info(f"Successfully parsed CSV file: {os.path.basename(file_path)}")
        return text
    except pd.errors.EmptyDataError:
        logging.warning(f"CSV file {os.path.basename(file_path)} is empty or has no columns")
        return "CSV file is empty or has no columns"
    except pd.errors.ParserError as e:
        logging.error(f"CSV parsing error for {file_path}: {e}")
        raise ValueError(f"Invalid CSV format: {e}")
    except UnicodeDecodeError as e:
        logging.error(f"Encoding error for CSV file {file_path}: {e}")
        raise UnicodeError(f"Unable to determine file encoding: {e}")
    except IOError as e:
        logging.error(f"IO error reading CSV file {file_path}: {e}")
        raise IOError(f"Error reading file: {e}")
    except Exception as e:
        logging.error(f"Error parsing CSV file {file_path}: {e}")
        raise

def parse_json_file(file_path):
    """
    Parses JSON files.

    Handles both JSON objects and arrays. For arrays of objects, attempts to convert
    to a tabular format for better readability.

    Args:
        file_path (str): Path to the JSON file

    Returns:
        str: Formatted string representation of the JSON content

    Raises:
        ValueError: If the file is not valid JSON
        UnicodeError: If the file encoding cannot be determined
        IOError: If there's an issue reading the file
        Exception: For other unexpected errors
    """
    try:
        logging.info(f"Parsing JSON file: {os.path.basename(file_path)}")

        # Check if file is empty
        if os.path.getsize(file_path) == 0:
            logging.warning(f"JSON file {os.path.basename(file_path)} is empty")
            return "Empty JSON file"

        # Try to load the JSON file
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            with open(file_path, 'r', encoding='latin1') as f:
                data = json.load(f)

        # Check if data is None or empty
        if data is None:
            logging.warning(f"JSON file {os.path.basename(file_path)} contains null value")
            return "JSON file contains null value"
        elif isinstance(data, dict) and not data:
            logging.warning(f"JSON file {os.path.basename(file_path)} contains empty object {{}}")
            return "JSON file contains empty object {}"
        elif isinstance(data, list) and not data:
            logging.warning(f"JSON file {os.path.basename(file_path)} contains empty array []")
            return "JSON file contains empty array []"

        # Format JSON data for readability
        if isinstance(data, list):
            text = f"JSON File Content (Array with {len(data)} items):\n\n"

            # If it's a list of objects, try to convert to a DataFrame for better formatting
            if data and isinstance(data[0], dict):
                try:
                    # Check if all objects have the same structure
                    keys = set(data[0].keys())
                    uniform_structure = all(set(item.keys()) == keys for item in data)

                    if uniform_structure:
                        df = pd.DataFrame(data)

                        # For very large DataFrames, limit the output
                        if len(df) > 100:
                            text += "Note: Showing first 50 and last 50 rows of a large JSON array.\n\n"
                            text += "First 50 rows:\n"
                            text += df.head(50).to_string()
                            text += "\n\n...\n\n"
                            text += "Last 50 rows:\n"
                            text += df.tail(50).to_string()
                        else:
                            text += df.to_string()
                    else:
                        # If objects have different structures, fall back to pretty-printed JSON
                        text += "Note: Array contains objects with different structures.\n\n"

                        # For very large arrays, limit the output
                        if len(data) > 20:
                            text += json.dumps(data[:10], indent=2)
                            text += "\n\n... (truncated) ...\n\n"
                            text += json.dumps(data[-10:], indent=2)
                        else:
                            text += json.dumps(data, indent=2)
                except Exception as df_error:
                    logging.warning(f"Failed to convert JSON array to DataFrame: {df_error}")
                    # Fall back to pretty-printed JSON

                    # For very large arrays, limit the output
                    if len(data) > 20:
                        text += "Note: Showing first 10 and last 10 items of a large JSON array.\n\n"
                        text += json.dumps(data[:10], indent=2)
                        text += "\n\n... (truncated) ...\n\n"
                        text += json.dumps(data[-10:], indent=2)
                    else:
                        text += json.dumps(data, indent=2)
            else:
                # Just pretty-print the JSON array

                # For very large arrays, limit the output
                if len(data) > 100:
                    text += "Note: Showing first 50 and last 50 items of a large JSON array.\n\n"
                    text += json.dumps(data[:50], indent=2)
                    text += "\n\n... (truncated) ...\n\n"
                    text += json.dumps(data[-50:], indent=2)
                else:
                    text += json.dumps(data, indent=2)
        else:
            # For objects/dictionaries
            text = "JSON File Content (Object):\n\n"

            # Get the number of keys at the top level
            num_keys = len(data.keys()) if isinstance(data, dict) else 0

            # Add metadata about the object
            if isinstance(data, dict):
                text += f"Object with {num_keys} top-level keys: {', '.join(data.keys())}\n\n"

            # Pretty-print the JSON object
            text += json.dumps(data, indent=2)

        logging.info(f"Successfully parsed JSON file: {os.path.basename(file_path)}")
        return text
    except json.JSONDecodeError as e:
        logging.error(f"Invalid JSON in file {file_path}: {e}")
        # Provide more helpful error message with line and column information
        raise ValueError(f"Invalid JSON format in file: {os.path.basename(file_path)} at line {e.lineno}, column {e.colno}: {e.msg}") from e
    except UnicodeDecodeError as e:
        logging.error(f"Encoding error for JSON file {file_path}: {e}")
        raise UnicodeError(f"Unable to determine file encoding: {e}")
    except IOError as e:
        logging.error(f"IO error reading JSON file {file_path}: {e}")
        raise IOError(f"Error reading file: {e}")
    except Exception as e:
        logging.error(f"Error parsing JSON file {file_path}: {e}")
        raise

def parse_html_file(file_path):
    """
    Parses HTML files using BeautifulSoup.

    Extracts text content from HTML files, removing scripts, styles, and other non-content elements.
    Also extracts the document title if available.

    Args:
        file_path (str): Path to the HTML file

    Returns:
        str: Extracted text content from the HTML file

    Raises:
        ValueError: If the file is not valid HTML
        UnicodeError: If the file encoding cannot be determined
        IOError: If there's an issue reading the file
        Exception: For other unexpected errors
    """
    try:
        logging.info(f"Parsing HTML file: {os.path.basename(file_path)}")

        # Check if file is empty
        if os.path.getsize(file_path) == 0:
            logging.warning(f"HTML file {os.path.basename(file_path)} is empty")
            return "Empty HTML file"

        # Try to read the file with different encodings if needed
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            with open(file_path, 'r', encoding='latin1', errors='replace') as f:
                html_content = f.read()
                logging.info(f"Used latin1 encoding for HTML file: {os.path.basename(file_path)}")

        # Check if content is too small to be valid HTML
        if len(html_content) < 10:  # Arbitrary small size
            logging.warning(f"HTML file {os.path.basename(file_path)} contains minimal content")
            return f"HTML file contains minimal content: {html_content}"

        # Parse HTML with BeautifulSoup
        try:
            soup = BeautifulSoup(html_content, 'lxml')
        except Exception as parser_error:
            # Fall back to html.parser if lxml fails
            logging.warning(f"lxml parser failed, falling back to html.parser: {parser_error}")
            soup = BeautifulSoup(html_content, 'html.parser')

        # Extract metadata
        metadata = {}

        # Extract title
        title = soup.title.string if soup.title else "No title"
        metadata["title"] = title

        # Extract meta description if available
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc and 'content' in meta_desc.attrs:
            metadata["description"] = meta_desc['content']

        # Extract meta keywords if available
        meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
        if meta_keywords and 'content' in meta_keywords.attrs:
            metadata["keywords"] = meta_keywords['content']

        # Extract headings for structure overview
        headings = []
        for h in soup.find_all(['h1', 'h2', 'h3']):
            if h.text.strip():
                headings.append(f"{h.name}: {h.text.strip()}")

        # Extract text content (remove script, style, and other non-content elements)
        for element in soup(['script', 'style', 'head', 'title', 'meta', 'link', 'noscript', 'iframe', 'svg']):
            element.extract()

        # Get text and clean up whitespace
        text = soup.get_text(separator='\n')
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        text = '\n'.join(lines)

        # Check if we extracted any meaningful content
        if not text:
            logging.warning(f"No text content extracted from HTML file: {os.path.basename(file_path)}")
            return f"HTML Document: {title}\n\nNo text content could be extracted from this HTML file."

        # Format the output
        result = f"HTML Document: {title}\n\n"

        # Add metadata if available
        if len(metadata) > 1:  # More than just title
            result += "Metadata:\n"
            for key, value in metadata.items():
                if key != "title":  # Title already shown
                    result += f"- {key}: {value}\n"
            result += "\n"

        # Add headings overview if available
        if headings:
            result += "Document Structure:\n"
            # Limit to first 10 headings if there are many
            if len(headings) > 10:
                for h in headings[:10]:
                    result += f"- {h}\n"
                result += f"- ... ({len(headings) - 10} more headings)\n"
            else:
                for h in headings:
                    result += f"- {h}\n"
            result += "\n"

        # Add the main content
        result += "Content:\n"

        # For very large content, limit the output
        if len(text) > 10000:
            # Show first 5000 and last 2000 characters
            result += text[:5000]
            result += "\n\n... (content truncated) ...\n\n"
            result += text[-2000:]
            logging.info(f"Truncated large HTML content for {os.path.basename(file_path)}")
        else:
            result += text

        logging.info(f"Successfully parsed HTML file: {os.path.basename(file_path)}")
        return result
    except UnicodeDecodeError as e:
        logging.error(f"Encoding error for HTML file {file_path}: {e}")
        raise UnicodeError(f"Unable to determine file encoding: {e}")
    except IOError as e:
        logging.error(f"IO error reading HTML file {file_path}: {e}")
        raise IOError(f"Error reading file: {e}")
    except Exception as e:
        logging.error(f"Error parsing HTML file {file_path}: {e}")
        raise

def parse_odt_file(file_path):
    """
    Parses ODT (OpenDocument Text) files using odfpy.

    Extracts text content from ODT files, preserving paragraph structure.

    Args:
        file_path (str): Path to the ODT file

    Returns:
        str: Extracted text content from the ODT file

    Raises:
        ValueError: If the file is not a valid ODT file
        IOError: If there's an issue reading the file
        Exception: For other unexpected errors
    """
    try:
        logging.info(f"Parsing ODT file: {os.path.basename(file_path)}")

        # Check if file is empty
        if os.path.getsize(file_path) == 0:
            logging.warning(f"ODT file {os.path.basename(file_path)} is empty")
            return "Empty ODT file"

        # Check if file is too small to be a valid ODT file
        if os.path.getsize(file_path) < 100:  # Arbitrary small size
            logging.warning(f"ODT file {os.path.basename(file_path)} is too small to be valid")
            return "File is too small to be a valid ODT document"

        try:
            # Load the ODT document
            doc = odf.opendocument.load(file_path)
        except Exception as load_error:
            logging.error(f"Failed to load ODT file {file_path}: {load_error}")
            raise ValueError(f"Invalid ODT format: {load_error}")

        # Extract metadata if available
        metadata = {}

        # Try to get document title
        try:
            metadata["title"] = doc.meta.getChildByName("title").firstChild.data
        except:
            metadata["title"] = "Untitled"

        # Try to get document creator
        try:
            metadata["creator"] = doc.meta.getChildByName("creator").firstChild.data
        except:
            pass

        # Try to get document creation date
        try:
            metadata["date"] = doc.meta.getChildByName("creation-date").firstChild.data
        except:
            pass

        # Extract all paragraphs
        paragraphs = doc.getElementsByType(P)

        # Check if document has any paragraphs
        if not paragraphs:
            logging.warning(f"ODT file {os.path.basename(file_path)} contains no paragraphs")
            return "ODT document contains no text content"

        # Combine text from all paragraphs
        text = ""
        for paragraph in paragraphs:
            # Get the text content of the paragraph by iterating through its children
            para_text = ""
            for node in paragraph.childNodes:
                if node.nodeType == node.TEXT_NODE:
                    para_text += node.data
                elif hasattr(node, 'childNodes'):
                    # Handle spans and other elements with text
                    for child in node.childNodes:
                        if child.nodeType == child.TEXT_NODE:
                            para_text += child.data

            if para_text:
                text += para_text + "\n"

        # Check if we extracted any meaningful content
        if not text.strip():
            logging.warning(f"No text content extracted from ODT file: {os.path.basename(file_path)}")
            return "ODT document contains no extractable text content"

        # Format the output
        result = ""

        # Add metadata if available
        if metadata:
            if "title" in metadata and metadata["title"] != "Untitled":
                result += f"Title: {metadata['title']}\n"
            if "creator" in metadata:
                result += f"Author: {metadata['creator']}\n"
            if "date" in metadata:
                result += f"Created: {metadata['date']}\n"
            if result:
                result += "\n"

        # Add the main content
        result += text

        # For very large content, limit the output
        if len(result) > 10000:
            # Show first 5000 and last 2000 characters
            truncated = result[:5000] + "\n\n... (content truncated) ...\n\n" + result[-2000:]
            logging.info(f"Truncated large ODT content for {os.path.basename(file_path)}")
            return truncated

        logging.info(f"Successfully parsed ODT file: {os.path.basename(file_path)}")
        return result
    except IOError as e:
        logging.error(f"IO error reading ODT file {file_path}: {e}")
        raise IOError(f"Error reading file: {e}")
    except ValueError as e:
        # Re-raise ValueError (invalid format)
        logging.error(f"Invalid ODT format in file {file_path}: {e}")
        raise
    except Exception as e:
        logging.error(f"Error parsing ODT file {file_path}: {e}")
        raise

# --- AI Processing Functions ---

def get_ai_table_processing(content, file_extension, model_name="gemini-1.5-flash-latest"):
    """
    Processes tabular data (Excel, CSV) using Gemini API and returns a structured table format.

    Args:
        content (str): The raw content extracted from the tabular file
        file_extension (str): The file extension (.xlsx, .csv, etc.) to determine the format
        model_name (str): The Gemini model to use

    Returns:
        str: Processed tabular data in a structured format with analysis
    """
    if not gemini_api_key:
        logging.error("Cannot call Gemini API: API key not configured.")
        return "[AI Processing Skipped: API Key Missing]"

    try:
        # Use gemini-1.5-flash-latest model
        model = genai.GenerativeModel(model_name)

        # Create a detailed prompt for tabular data processing
        prompt = f"""Zpracuj následující tabulková data s maximální přesností (95%+).

        Typ souboru: {file_extension}

        Tabulková data:
        ---
        {content}
        ---

        Proveď následující:
        1. Zachovej tabulkovou strukturu dat v odpovědi - používej formátování pomocí | a - pro vytvoření markdown tabulek
        2. Identifikuj hlavičky sloupců a datové typy
        3. Poskytni stručnou analýzu dat (trendy, vzory, anomálie)
        4. Pokud jsou data příliš rozsáhlá, zaměř se na nejdůležitější části, ale zachovej tabulkovou strukturu
        5. Pokud data obsahují číselné hodnoty, uveď základní statistiky (min, max, průměr, apod.)
        6. Nezjednodušuj ani nezobecňuj obsah - zachovej přesné hodnoty

        Tvůj výstup by měl obsahovat:
        1. Strukturovanou tabulku v markdown formátu
        2. Analýzu dat
        3. Případné statistiky

        Formátuj výstup tak, aby byl přehledný a zachoval tabulkovou strukturu.
        """

        # Configure generation parameters
        generation_config = {
            "temperature": 0.1,  # Very low temperature for precise output
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 8192,
        }

        logging.info(f"Sending request to Gemini API for tabular data processing")
        response = model.generate_content(prompt, generation_config=generation_config)

        if not response.text:
            logging.warning(f"Gemini API returned empty response for tabular data processing")
            return "[AI did not provide any processed tabular data]"

        logging.info(f"Successfully processed tabular data")
        return response.text
    except Exception as e:
        logging.error(f"Error calling Gemini API for tabular data ({model_name}): {e}")
        return f"[AI Processing Error for tabular data: {e}]"

def get_ai_image_description(file_path, model_name="gemini-1.5-flash-latest"):
    """Generates a detailed description for an image file using Gemini multimodal API."""
    if not gemini_api_key:
        logging.error("Cannot call Gemini API: API key not configured.")
        return "[AI Processing Skipped: API Key Missing]"

    try:
        logging.info(f"Loading image for AI description: {file_path}")
        img = Image.open(file_path)

        # Use gemini-1.5-flash-latest which supports multimodal input
        model = genai.GenerativeModel(model_name)

        # Prepare prompt with image and detailed instruction for high accuracy
        prompt_parts = [
            """Popiš prosím tento obrázek s maximální přesností (95%+).
            Uveď všechny důležité detaily, které vidíš.
            Nezjednodušuj ani nezobecňuj obsah.
            Popiš vše, co je na obrázku viditelné, včetně textu, barev, objektů, osob, pozadí a dalších prvků.
            Pokud je na obrázku text, přepiš ho přesně.
            Tvůj popis by měl být vyčerpávající a zachytit téměř veškerý obsah obrázku.""", # Detailed instruction in Czech
            img,
        ]
        logging.info(f"Sending request to Gemini API for detailed image description: {os.path.basename(file_path)}")
        response = model.generate_content(prompt_parts)
        logging.debug(f"Raw Gemini API response for image: {response}") # Log raw response for debugging

        # Check for response content and potential errors/blocks
        try:
            # Accessing response.text directly might raise if blocked, etc.
            description = response.text
            if not description:
                 logging.warning(f"Gemini API returned empty description for image: {os.path.basename(file_path)}")
                 description = "[AI did not provide a description]"
            else:
                 logging.info(f"Successfully generated detailed image description for {os.path.basename(file_path)}")
        except ValueError as ve:
            # Handle cases where accessing .text fails (e.g., blocked prompt)
            logging.error(f"Could not extract text from Gemini response for image {os.path.basename(file_path)}. Possible block? Response: {response}. Error: {ve}")
            description = f"[AI Response Error: Could not extract text - {ve}]"
        except Exception as resp_err:
            # Catch other potential issues with the response object
            logging.error(f"Unexpected error processing Gemini response for image {os.path.basename(file_path)}: {resp_err}")
            description = f"[AI Response Processing Error: {resp_err}]"

        return description
    except FileNotFoundError:
        logging.error(f"Image file not found during AI processing: {file_path}")
        raise # Re-raise to be caught by process_file
    except Exception as e:
        # Log more details from the exception if possible
        logging.error(f"Error calling Gemini API for image ({model_name}): {e}", exc_info=True) # Add traceback
        # Attempt to get more specific error details if it's a google API error
        error_details = str(e)
        if hasattr(e, 'details'):
             error_details = f"{e} - Details: {e.details()}"
        elif hasattr(e, 'message'):
             error_details = f"{e} - Message: {e.message}"
        return f"[AI Image Processing Error: {error_details}]"


def get_ai_summary(content, model_name="gemini-1.5-flash-latest", max_input_length=12000):
    """
    Generates a detailed content extraction with high accuracy using Gemini API.

    For large inputs, the content is truncated to avoid API limits.

    Args:
        content (str): The text content to summarize
        model_name (str): The Gemini model to use
        max_input_length (int): Maximum length of input text to send to API

    Returns:
        str: Processed content or error message
    """
    if not gemini_api_key:
        logging.error("Cannot call Gemini API: API key not configured.")
        return "[AI Processing Skipped: API Key Missing]"

    try:
        # Truncate content if it's too long to avoid API limits
        original_length = len(content)
        if original_length > max_input_length:
            logging.info(f"Content length ({original_length} chars) exceeds max input length ({max_input_length}). Truncating...")
            # Take first 70% and last 30% of allowed length to capture beginning and end
            first_part_size = int(max_input_length * 0.7)
            last_part_size = max_input_length - first_part_size
            truncated_content = content[:first_part_size] + "\n\n[...obsah zkrácen...]\n\n" + content[-last_part_size:]
            content = truncated_content
            logging.info(f"Content truncated to {len(content)} characters")

        # Use gemini-1.5-flash-latest as requested (efficient model)
        model = genai.GenerativeModel(model_name)

        # Detailed prompt for high accuracy content extraction
        prompt = f"""Zpracuj následující text s maximální přesností (95%+).
        Zachovej co nejvíce původního obsahu a informací.
        Nezjednodušuj ani nezobecňuj obsah.
        Pokud text obsahuje specifické údaje, čísla, data, jména nebo technické informace, zachovej je přesně.
        Tvůj výstup by měl obsahovat téměř veškerý důležitý obsah původního textu.
        Pokud je text příliš dlouhý, zaměř se na nejdůležitější části, ale stále zachovej vysokou míru detailů.

        Text k zpracování:
        ---
        {content}
        ---

        Detailní zpracování obsahu:"""

        # Configure generation parameters for more detailed output
        generation_config = {
            "temperature": 0.2,  # Lower temperature for more precise output
            "top_p": 0.95,       # Higher top_p for more comprehensive coverage
            "top_k": 40,         # Higher top_k for more diverse vocabulary when needed
            "max_output_tokens": 4096,  # Reduced from 8192 to avoid timeouts
        }

        logging.info(f"Sending request to Gemini API for detailed content extraction")

        # Add retry logic for API calls
        max_retries = 2
        retry_delay = 2  # seconds

        for attempt in range(max_retries + 1):
            try:
                response = model.generate_content(prompt, generation_config=generation_config)
                break  # Success, exit retry loop
            except Exception as retry_error:
                if attempt < max_retries:
                    logging.warning(f"API call attempt {attempt+1} failed: {retry_error}. Retrying in {retry_delay} seconds...")
                    import time
                    time.sleep(retry_delay)
                    retry_delay *= 2  # Exponential backoff
                else:
                    # Last attempt failed, re-raise the exception
                    raise

        if not response.text:
            logging.warning(f"Gemini API returned empty response for content extraction")
            return "[AI did not provide any content]"

        logging.info(f"Successfully generated detailed content extraction")

        # If content was truncated, add a note
        if original_length > max_input_length:
            return f"[Poznámka: Původní dokument byl příliš dlouhý ({original_length} znaků), byl zpracován pouze částečný obsah.]\n\n{response.text}"
        else:
            return response.text

    except Exception as e:
        logging.error(f"Error calling Gemini API ({model_name}): {e}")
        # Consider returning specific error messages based on API response if available
        return f"[AI Processing Error: {e}]"

# --- Main Processing Function ---

def process_file(file_path, max_file_size_mb=50):
    """
    Identifies file type, parses content, and optionally calls AI for processing.

    This is the main entry point for file conversion. It automatically detects the file type
    based on the file extension and calls the appropriate parsing function. The extracted
    content is then processed using AI (if applicable) to generate a summary or description.

    Supported file formats:
    - Text Documents: .txt, .pdf, .docx, .doc, .odt, .html, .htm
    - Structured Data: .xlsx, .xls, .csv, .json
    - Images: .png, .jpg, .jpeg, .webp, .gif, .bmp

    Args:
        file_path (str): Path to the file to process
        max_file_size_mb (int, optional): Maximum file size in megabytes. Defaults to 50MB.

    Returns:
        dict: A dictionary with the following keys:
            - content (str or None): The processed content if successful, None otherwise
            - error (str or None): Error message if an error occurred, None otherwise

    Example:
        >>> result = process_file("document.pdf")
        >>> if result["error"]:
        ...     print(f"Error: {result['error']}")
        ... else:
        ...     print(f"Content: {result['content']}")
    """
    result = {"content": None, "error": None}
    try:
        file_path_obj = Path(file_path)
        if not file_path_obj.is_file():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Check file size
        file_size_mb = file_path_obj.stat().st_size / (1024 * 1024)  # Convert bytes to MB
        if file_size_mb > max_file_size_mb:
            raise ValueError(f"File size ({file_size_mb:.2f}MB) exceeds maximum allowed size ({max_file_size_mb}MB)")

        file_extension = file_path_obj.suffix.lower()
        processed_content = None # Initialize processed_content

        logging.info(f"Processing file: {file_path} (Extension: {file_extension})")

        # 1. Parse Content / Process based on file type
        if file_extension == '.txt':
            raw_content = parse_text_file(file_path)
            if raw_content:
                # Check if the content is an email correspondence
                is_email, email_json = process_email_correspondence(raw_content)
                if is_email:
                    # If it's an email, return the structured JSON
                    logging.info(f"Detected email correspondence in {os.path.basename(file_path)}")
                    processed_content = email_json
                else:
                    # Otherwise, process as regular text
                    processed_content = get_ai_summary(raw_content)
            else: # Handle parsing error case
                 result["error"] = f"Failed to parse text content from {os.path.basename(file_path)}"
        elif file_extension == '.pdf':
            try:
                # Get total pages in PDF to provide better information
                try:
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        total_pages = len(reader.pages)
                        logging.info(f"PDF has {total_pages} pages in total")
                except Exception as page_count_err:
                    logging.warning(f"Could not determine total pages in PDF: {page_count_err}")
                    total_pages = "unknown"

                # Set max pages based on total pages to avoid processing too many
                max_pages_to_process = 5  # Default max pages

                # Try to parse PDF with OCR fallback enabled and limited pages
                raw_content = parse_pdf_file(file_path, max_pages=max_pages_to_process, try_ocr=True)

                if raw_content and raw_content.strip(): # Check if parsing returned non-empty content
                    # Add note about limited pages if needed
                    if isinstance(total_pages, int) and total_pages > max_pages_to_process:
                        raw_content += f"\n\n[Poznámka: Dokument obsahuje celkem {total_pages} stránek, ale bylo zpracováno pouze prvních {max_pages_to_process} stránek.]"

                    # Check if the content is an email correspondence
                    is_email, email_json = process_email_correspondence(raw_content)
                    if is_email:
                        # If it's an email, return the structured JSON
                        logging.info(f"Detected email correspondence in PDF file: {os.path.basename(file_path)}")
                        processed_content = email_json
                    else:
                        # Otherwise, process as regular text
                        processed_content = get_ai_summary(raw_content)
                else: # Handle empty content case
                    # Try to use AI image description as a fallback for PDFs that might be image-based
                    logging.info(f"PDF text extraction returned empty result, trying AI image description for first page: {os.path.basename(file_path)}")
                    try:
                        # Convert first page to image and use image description
                        if OCR_AVAILABLE:
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=True) as temp_img:
                                # Convert first page to image
                                images = pdf2image.convert_from_path(file_path, first_page=1, last_page=1)
                                if images:
                                    images[0].save(temp_img.name)
                                    # Use AI image description
                                    processed_content = get_ai_image_description(temp_img.name)

                                    # Add note about limited processing
                                    if isinstance(total_pages, int) and total_pages > 1:
                                        processed_content = f"[PDF zpracován jako obrázek - pouze první stránka z celkem {total_pages} stránek] {processed_content}"
                                    else:
                                        processed_content = f"[PDF zpracován jako obrázek] {processed_content}"
                                else:
                                    result["error"] = f"Failed to extract content from PDF: {os.path.basename(file_path)}"
                        else:
                            result["error"] = f"Failed to extract text content from PDF and OCR is not available: {os.path.basename(file_path)}"
                    except Exception as img_err:
                        logging.error(f"Error during PDF image fallback: {img_err}")
                        result["error"] = f"Failed to extract any content from PDF: {os.path.basename(file_path)}"
            except Exception as pdf_err:
                logging.error(f"Error processing PDF file: {pdf_err}")
                result["error"] = f"Error processing PDF file: {str(pdf_err)}"
        elif file_extension == '.docx':
            # Nejprve zkontrolujeme, zda soubor může obsahovat emailovou korespondenci
            # Otevřeme soubor a zkontrolujeme prvních 50 odstavců
            try:
                doc = docx.Document(file_path)
                email_indicators = ['Od:', 'Odesílatel:', 'From:', 'Komu:', 'To:', 'Předmět:', 'Subject:']
                email_indicator_count = 0

                # Zkontrolujeme prvních 50 odstavců (nebo méně, pokud jich je méně)
                for i in range(min(50, len(doc.paragraphs))):
                    para_text = doc.paragraphs[i].text
                    for indicator in email_indicators:
                        if indicator in para_text:
                            email_indicator_count += 1

                # Pokud najdeme alespoň 3 indikátory, považujeme to za emailovou korespondenci
                is_likely_email = email_indicator_count >= 3

                if is_likely_email:
                    logging.info(f"Document appears to be email correspondence, processing all pages")
                    # Pro emailovou korespondenci zpracujeme celý dokument
                    raw_content = parse_docx_file(file_path, max_pages=None)
                else:
                    # Pro běžné dokumenty omezíme počet stránek
                    raw_content = parse_docx_file(file_path, max_pages=25)
            except Exception as e:
                logging.error(f"Error pre-checking DOCX file: {e}")
                # Pokud se nepodaří provést předběžnou kontrolu, zpracujeme dokument standardně
                raw_content = parse_docx_file(file_path)

            if raw_content: # Check if parsing returned content
                # Check if the content is an email correspondence
                is_email, email_json = process_email_correspondence(raw_content)
                if is_email:
                    # If it's an email, return the structured JSON
                    logging.info(f"Detected email correspondence in DOCX file: {os.path.basename(file_path)}")
                    processed_content = email_json
                else:
                    # Otherwise, process as regular text
                    processed_content = get_ai_summary(raw_content)
            else: # Handle parsing error case
                 result["error"] = f"Failed to extract text content from DOCX: {os.path.basename(file_path)}"
                 logging.warning(result["error"])
        elif file_extension == '.doc':
             raw_content = parse_doc_file(file_path)
             if raw_content:
                 # Check if the content is an error/installation message
                 if raw_content.strip().startswith("Antiword is not installed") or raw_content.strip().startswith("Error processing .doc file"):
                     # This is an installation message or error, not actual content
                     result["content"] = raw_content  # Set as content so it's displayed to the user
                     result["error"] = "DOC file processing requires antiword"
                     logging.warning(f"DOC file processing requires antiword: {os.path.basename(file_path)}")
                 else:
                     # Check if the content is an email correspondence
                     is_email, email_json = process_email_correspondence(raw_content)
                     if is_email:
                         # If it's an email, return the structured JSON
                         logging.info(f"Detected email correspondence in DOC file: {os.path.basename(file_path)}")
                         processed_content = email_json
                     else:
                         # Normal content, process with AI
                         processed_content = get_ai_summary(raw_content)
             else:
                 result["error"] = f"Failed to extract text content from DOC: {os.path.basename(file_path)}"
                 logging.warning(result["error"])
        elif file_extension in ['.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp']: # Common image types
             # For images, we directly get the description, no intermediate raw_content needed for summary
             processed_content = get_ai_image_description(file_path)
        elif file_extension == '.xlsx' or file_extension == '.xls':
            raw_content = parse_xlsx_file(file_path)
            if raw_content:
                # Use the specialized table processing function for Excel files
                processed_content = get_ai_table_processing(raw_content, file_extension)
            else:
                result["error"] = f"Failed to extract content from Excel file: {os.path.basename(file_path)}"
                logging.warning(result["error"])
        elif file_extension == '.csv':
            raw_content = parse_csv_file(file_path)
            if raw_content:
                # Use the specialized table processing function for CSV files
                processed_content = get_ai_table_processing(raw_content, file_extension)
            else:
                result["error"] = f"Failed to extract content from CSV file: {os.path.basename(file_path)}"
                logging.warning(result["error"])
        elif file_extension == '.json':
            raw_content = parse_json_file(file_path)
            if raw_content:
                processed_content = get_ai_summary(raw_content)
            else:
                result["error"] = f"Failed to extract content from JSON file: {os.path.basename(file_path)}"
                logging.warning(result["error"])
        elif file_extension == '.html' or file_extension == '.htm':
            raw_content = parse_html_file(file_path)
            if raw_content:
                processed_content = get_ai_summary(raw_content)
            else:
                result["error"] = f"Failed to extract content from HTML file: {os.path.basename(file_path)}"
                logging.warning(result["error"])
        elif file_extension == '.odt':
            raw_content = parse_odt_file(file_path)
            if raw_content:
                processed_content = get_ai_summary(raw_content)
            else:
                result["error"] = f"Failed to extract content from ODT file: {os.path.basename(file_path)}"
                logging.warning(result["error"])
        else:
            result["error"] = f"Unsupported file type: {file_extension}"
            logging.warning(result["error"])
            # No need to return here, let the finally block handle it

        # 2. Final check and assignment to result dict
        if result["error"] is None: # Only proceed if no parsing/type error occurred yet
            if processed_content is None:
                 # This might happen if parsing succeeded but AI processing failed internally without setting processed_content
                 result["error"] = f"Failed to generate processed content from {os.path.basename(file_path)}"
                 logging.warning(result["error"])
            elif isinstance(processed_content, str) and processed_content.startswith("[AI"): # Check if AI returned an error/placeholder string
                 if "Error" in processed_content or "Skipped" in processed_content:
                     result["error"] = processed_content # Log AI error
                     if result["content"] is None:  # Only set to None if not already set (for DOC files we might have set content)
                         result["content"] = None # No successful content
                     logging.warning(f"AI processing issue for {os.path.basename(file_path)}: {processed_content}")
                 else:
                     # It might be a placeholder like "[AI did not provide a description]"
                     result["content"] = processed_content # Keep the placeholder as content
            else:
                 # Processing seems successful
                 result["content"] = processed_content
                 logging.info(f"Successfully processed file: {os.path.basename(file_path)}")
        elif result["content"] is not None:
            # Special case: we have both error and content (e.g., for DOC files without antiword)
            logging.info(f"File processed with warning: {os.path.basename(file_path)}")

    except FileNotFoundError as e:
        result["error"] = str(e)
        logging.error(result["error"])
    except Exception as e:
        result["error"] = f"An unexpected error occurred during processing: {e}"
        logging.exception("Unexpected error in process_file:") # Log full traceback
    finally:
        return result

# Example Usage (for testing)
if __name__ == '__main__':
    # Create dummy files for testing
    if not os.path.exists("test_files"):
        os.makedirs("test_files")
    with open("test_files/sample.txt", "w") as f:
        f.write("Toto je jednoduchý textový soubor pro testování.")
    # You would need actual PDF/DOCX files for those tests

    print("--- Testing Text File ---")
    test_result_txt = process_file("test_files/sample.txt")
    print(f"Result: {test_result_txt}")

    # print("\n--- Testing PDF File (Placeholder) ---")
    # test_result_pdf = process_file("test_files/sample.pdf") # Needs a real PDF
    # print(f"Result: {test_result_pdf}")

    # print("\n--- Testing DOCX File (Placeholder) ---")
    # test_result_docx = process_file("test_files/sample.docx") # Needs a real DOCX
    # print(f"Result: {test_result_docx}")

    print("\n--- Testing Non-existent File ---")
    test_result_missing = process_file("non_existent_file.xyz")
    print(f"Result: {test_result_missing}")

    print("\n--- Testing Unsupported File ---")
    with open("test_files/sample.xyz", "w") as f:
        f.write("some data")
    test_result_unsupported = process_file("test_files/sample.xyz")
    print(f"Result: {test_result_unsupported}")

    # Create and test JSON file
    print("\n--- Testing JSON File ---")
    with open("test_files/sample.json", "w") as f:
        f.write('{"name": "Test User", "age": 30, "email": "test@example.com"}')
    test_result_json = process_file("test_files/sample.json")
    print(f"Result: {test_result_json}")

    # Create and test CSV file
    print("\n--- Testing CSV File ---")
    with open("test_files/sample.csv", "w") as f:
        f.write("name,age,email\nTest User,30,test@example.com\nAnother User,25,another@example.com")
    test_result_csv = process_file("test_files/sample.csv")
    print(f"Result: {test_result_csv}")

    # Create and test Excel file if pandas is available
    try:
        print("\n--- Testing Excel File ---")
        # Create a simple DataFrame and save as Excel
        df = pd.DataFrame({
            'Name': ['Test User', 'Another User'],
            'Age': [30, 25],
            'Email': ['test@example.com', 'another@example.com']
        })
        df.to_excel("test_files/sample.xlsx", index=False)
        test_result_xlsx = process_file("test_files/sample.xlsx")
        print(f"Result: {test_result_xlsx}")
    except Exception as e:
        print(f"Error during Excel test: {e}")

    # print("\n--- Testing Image File (Placeholder) ---")
    # # Needs a real image file and API key configured
    # # Create a dummy png for structure test if needed
    # try:
    #     # from PIL import Image # Already imported at top
    #     dummy_img = Image.new('RGB', (60, 30), color = 'red')
    #     dummy_img.save("test_files/sample.png")
    #     test_result_png = process_file("test_files/sample.png")
    #     print(f"Result (PNG): {test_result_png}")
    # except ImportError:
    #     print("Pillow not installed, skipping image test.")
    # except Exception as e:
    #     print(f"Error during image test setup/run: {e}")
